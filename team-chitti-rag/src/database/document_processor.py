"""
Document Processor - Handle various file formats and web links
"""

import os
import json
import csv
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import requests
from urllib.parse import urlparse
import time
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a processed document chunk"""
    content: str
    metadata: Dict[str, Any]
    chunk_id: str
    source_file: str
    file_type: str


class DocumentProcessor:
    """Process various document formats and web content"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize document processor
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Supported file extensions
        self.supported_extensions = {
            '.txt', '.md', '.csv', '.json', 
            '.pdf', '.docx', '.doc', '.html', '.htm'
        }
        
        logger.info("🔧 Document processor initialized")
    
    def process_directory(self, directory_path: str) -> List[DocumentChunk]:
        """Process all supported files in a directory
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of processed document chunks
        """
        directory = Path(directory_path)
        if not directory.exists():
            logger.error(f"❌ Directory not found: {directory_path}")
            return []
        
        all_chunks = []
        
        # Process all files recursively
        for file_path in directory.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    chunks = self.process_file(str(file_path))
                    all_chunks.extend(chunks)
                    logger.info(f"✅ Processed: {file_path.name} ({len(chunks)} chunks)")
                except Exception as e:
                    logger.error(f"❌ Error processing {file_path}: {e}")
        
        logger.info(f"📚 Total processed chunks: {len(all_chunks)}")
        return all_chunks
    
    def process_file(self, file_path: str) -> List[DocumentChunk]:
        """Process a single file based on its extension
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            List of document chunks
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension in ['.txt', '.md']:
            return self._process_text_file(file_path)
        elif extension == '.csv':
            return self._process_csv_file(file_path)
        elif extension == '.json':
            return self._process_json_file(file_path)
        elif extension == '.pdf':
            return self._process_pdf_file(file_path)
        elif extension in ['.docx', '.doc']:
            return self._process_docx_file(file_path)
        elif extension in ['.html', '.htm']:
            return self._process_html_file(file_path)
        else:
            logger.warning(f"⚠️  Unsupported file type: {extension}")
            return []
    
    def _process_text_file(self, file_path: Path) -> List[DocumentChunk]:
        """Process text/markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            chunks = self._chunk_text(content)
            
            return [
                DocumentChunk(
                    content=chunk,
                    metadata={
                        "file_name": file_path.name,
                        "file_size": file_path.stat().st_size,
                        "file_type": "text",
                        "encoding": "utf-8"
                    },
                    chunk_id=f"{file_path.stem}_chunk_{i}",
                    source_file=str(file_path),
                    file_type="text"
                )
                for i, chunk in enumerate(chunks)
            ]
            
        except Exception as e:
            logger.error(f"❌ Error processing text file {file_path}: {e}")
            return []
    
    def _process_csv_file(self, file_path: Path) -> List[DocumentChunk]:
        """Process CSV files"""
        try:
            chunks = []
            
            with open(file_path, 'r', encoding='utf-8') as f:
                # Read CSV and convert to readable text
                csv_reader = csv.DictReader(f)
                headers = csv_reader.fieldnames
                
                # Create header chunk
                header_text = f"CSV File: {file_path.name}\nColumns: {', '.join(headers)}\n\n"
                
                # Process rows in batches
                rows_text = ""
                row_count = 0
                
                for row in csv_reader:
                    row_text = " | ".join([f"{k}: {v}" for k, v in row.items()])
                    rows_text += f"Row {row_count + 1}: {row_text}\n"
                    row_count += 1
                    
                    # Create chunk when reaching limit
                    if len(rows_text) >= self.chunk_size:
                        full_text = header_text + rows_text
                        chunks.append(
                            DocumentChunk(
                                content=full_text,
                                metadata={
                                    "file_name": file_path.name,
                                    "file_type": "csv",
                                    "headers": headers,
                                    "rows_in_chunk": rows_text.count('\n'),
                                    "total_rows": row_count
                                },
                                chunk_id=f"{file_path.stem}_csv_chunk_{len(chunks)}",
                                source_file=str(file_path),
                                file_type="csv"
                            )
                        )
                        rows_text = ""
                
                # Handle remaining rows
                if rows_text:
                    full_text = header_text + rows_text
                    chunks.append(
                        DocumentChunk(
                            content=full_text,
                            metadata={
                                "file_name": file_path.name,
                                "file_type": "csv",
                                "headers": headers,
                                "rows_in_chunk": rows_text.count('\n'),
                                "total_rows": row_count
                            },
                            chunk_id=f"{file_path.stem}_csv_chunk_{len(chunks)}",
                            source_file=str(file_path),
                            file_type="csv"
                        )
                    )
            
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Error processing CSV file {file_path}: {e}")
            return []
    
    def _process_json_file(self, file_path: Path) -> List[DocumentChunk]:
        """Process JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to readable text
            def json_to_text(obj, level=0):
                text = ""
                indent = "  " * level
                
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            text += f"{indent}{key}:\n"
                            text += json_to_text(value, level + 1)
                        else:
                            text += f"{indent}{key}: {value}\n"
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        text += f"{indent}Item {i + 1}:\n"
                        text += json_to_text(item, level + 1)
                else:
                    text += f"{indent}{obj}\n"
                
                return text
            
            readable_text = f"JSON File: {file_path.name}\n\n"
            readable_text += json_to_text(data)
            
            chunks = self._chunk_text(readable_text)
            
            return [
                DocumentChunk(
                    content=chunk,
                    metadata={
                        "file_name": file_path.name,
                        "file_type": "json",
                        "json_structure": type(data).__name__,
                        "total_keys": len(data) if isinstance(data, dict) else None
                    },
                    chunk_id=f"{file_path.stem}_json_chunk_{i}",
                    source_file=str(file_path),
                    file_type="json"
                )
                for i, chunk in enumerate(chunks)
            ]
            
        except Exception as e:
            logger.error(f"❌ Error processing JSON file {file_path}: {e}")
            return []
    
    def _process_pdf_file(self, file_path: Path) -> List[DocumentChunk]:
        """Process PDF files"""
        try:
            # Try to import PyPDF2 or pypdf
            try:
                import pypdf
                
                with open(file_path, 'rb') as f:
                    pdf_reader = pypdf.PdfReader(f)
                    text = ""
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                chunks = self._chunk_text(text)
                
                return [
                    DocumentChunk(
                        content=chunk,
                        metadata={
                            "file_name": file_path.name,
                            "file_type": "pdf",
                            "total_pages": len(pdf_reader.pages),
                            "extracted_text_length": len(text)
                        },
                        chunk_id=f"{file_path.stem}_pdf_chunk_{i}",
                        source_file=str(file_path),
                        file_type="pdf"
                    )
                    for i, chunk in enumerate(chunks)
                ]
                
            except ImportError:
                logger.warning("⚠️  PyPDF library not installed. Skipping PDF processing.")
                return []
                
        except Exception as e:
            logger.error(f"❌ Error processing PDF file {file_path}: {e}")
            return []
    
    def _process_docx_file(self, file_path: Path) -> List[DocumentChunk]:
        """Process DOCX files"""
        try:
            # Try to import python-docx
            try:
                from docx import Document
                
                doc = Document(file_path)
                text = ""
                
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                chunks = self._chunk_text(text)
                
                return [
                    DocumentChunk(
                        content=chunk,
                        metadata={
                            "file_name": file_path.name,
                            "file_type": "docx",
                            "total_paragraphs": len(doc.paragraphs),
                            "extracted_text_length": len(text)
                        },
                        chunk_id=f"{file_path.stem}_docx_chunk_{i}",
                        source_file=str(file_path),
                        file_type="docx"
                    )
                    for i, chunk in enumerate(chunks)
                ]
                
            except ImportError:
                logger.warning("⚠️  python-docx library not installed. Skipping DOCX processing.")
                return []
                
        except Exception as e:
            logger.error(f"❌ Error processing DOCX file {file_path}: {e}")
            return []
    
    def _process_html_file(self, file_path: Path) -> List[DocumentChunk]:
        """Process HTML files"""
        try:
            # Try to import BeautifulSoup
            try:
                from bs4 import BeautifulSoup
                
                with open(file_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                
                soup = BeautifulSoup(html_content, 'html.parser')
                text = soup.get_text()
                
                chunks = self._chunk_text(text)
                
                return [
                    DocumentChunk(
                        content=chunk,
                        metadata={
                            "file_name": file_path.name,
                            "file_type": "html",
                            "title": soup.title.string if soup.title else None,
                            "extracted_text_length": len(text)
                        },
                        chunk_id=f"{file_path.stem}_html_chunk_{i}",
                        source_file=str(file_path),
                        file_type="html"
                    )
                    for i, chunk in enumerate(chunks)
                ]
                
            except ImportError:
                logger.warning("⚠️  BeautifulSoup library not installed. Skipping HTML processing.")
                return []
                
        except Exception as e:
            logger.error(f"❌ Error processing HTML file {file_path}: {e}")
            return []
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                for boundary in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                    last_boundary = text.rfind(boundary, start, end)
                    if last_boundary != -1:
                        end = last_boundary + len(boundary)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks
    
    def process_web_links(self, links_file: str) -> List[DocumentChunk]:
        """Process web links from JSON file
        
        Args:
            links_file: Path to JSON file containing links
            
        Returns:
            List of processed document chunks from web content
        """
        try:
            with open(links_file, 'r', encoding='utf-8') as f:
                links_data = json.load(f)
            
            all_chunks = []
            
            for item in links_data:
                url = item.get('url', '')
                title = item.get('title', '')
                description = item.get('description', '')
                
                try:
                    chunks = self._download_and_process_url(url, title, description)
                    all_chunks.extend(chunks)
                    logger.info(f"✅ Processed URL: {url} ({len(chunks)} chunks)")
                    
                    # Rate limiting to be respectful
                    time.sleep(1)
                    
                except Exception as e:
                    logger.error(f"❌ Error processing URL {url}: {e}")
            
            return all_chunks
            
        except Exception as e:
            logger.error(f"❌ Error processing links file {links_file}: {e}")
            return []
    
    def _download_and_process_url(self, url: str, title: str, description: str) -> List[DocumentChunk]:
        """Download and process content from URL"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Try to parse as HTML
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Extract title if not provided
                if not title and soup.title:
                    title = soup.title.string.strip()
                
                # Extract main content
                text = soup.get_text()
                
                # Clean up text
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                text = '\n'.join(lines)
                
                # Add metadata header
                header = f"Title: {title}\nURL: {url}\nDescription: {description}\n\n"
                full_text = header + text
                
                chunks = self._chunk_text(full_text)
                
                return [
                    DocumentChunk(
                        content=chunk,
                        metadata={
                            "source_url": url,
                            "title": title,
                            "description": description,
                            "file_type": "web_content",
                            "domain": urlparse(url).netloc,
                            "downloaded_at": time.strftime("%Y-%m-%d %H:%M:%S")
                        },
                        chunk_id=f"web_{urlparse(url).netloc}_{hash(url)}_chunk_{i}",
                        source_file=url,
                        file_type="web"
                    )
                    for i, chunk in enumerate(chunks)
                ]
                
            except ImportError:
                logger.warning("⚠️  BeautifulSoup not available. Processing as plain text.")
                text = response.text
                chunks = self._chunk_text(text)
                
                return [
                    DocumentChunk(
                        content=chunk,
                        metadata={
                            "source_url": url,
                            "title": title,
                            "description": description,
                            "file_type": "web_content_plain",
                            "domain": urlparse(url).netloc
                        },
                        chunk_id=f"web_{hash(url)}_chunk_{i}",
                        source_file=url,
                        file_type="web"
                    )
                    for i, chunk in enumerate(chunks)
                ]
                
        except Exception as e:
            logger.error(f"❌ Error downloading URL {url}: {e}")
            return []
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return sorted(list(self.supported_extensions))
    
    def get_processing_stats(self, chunks: List[DocumentChunk]) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        if not chunks:
            return {"total_chunks": 0}
        
        file_types = {}
        total_content_length = 0
        sources = set()
        
        for chunk in chunks:
            file_type = chunk.file_type
            file_types[file_type] = file_types.get(file_type, 0) + 1
            total_content_length += len(chunk.content)
            sources.add(chunk.source_file)
        
        return {
            "total_chunks": len(chunks),
            "unique_sources": len(sources),
            "file_types": file_types,
            "total_content_length": total_content_length,
            "average_chunk_size": total_content_length / len(chunks) if chunks else 0
        }